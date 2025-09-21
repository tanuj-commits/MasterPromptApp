import streamlit as st
from openai import OpenAI
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

# ------------------ CONFIG ------------------
st.set_page_config(page_title="Talent Acquisition Playbook Generator", layout="wide")
st.title("📘 Talent Acquisition Playbook Generator (Upload JD)")

st.markdown("""
Upload a Job Description (PDF or Word), select city tier and role/department, and generate a **consulting-style TA playbook**.
""")

# ------------------ API KEY INPUT ------------------
st.subheader("Step 0: Enter OpenAI API Key")
api_key = st.text_input(
    "Enter your OpenAI API Key:",
    type="password",
    help="You can find this key in your OpenAI account."
)

client = None
if api_key:
    try:
        client = OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {e}")

# ------------------ JD UPLOAD ------------------
st.subheader("Step 1: Upload JD (PDF or Word)")
uploaded_file = st.file_uploader(
    "Choose a file",
    type=["pdf", "docx"]
)

jd_text = ""
if uploaded_file is not None:
    if uploaded_file.type == "application/pdf":
        reader = PdfReader(uploaded_file)
        jd_text = ""
        for page in reader.pages:
            jd_text += page.extract_text() + "\n"
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        jd_text = "\n".join([para.text for para in doc.paragraphs])
    else:
        st.error("Unsupported file type. Please upload PDF or Word (.docx).")

# Display JD preview
if jd_text:
    st.text_area("JD Preview", jd_text, height=200)

# ------------------ ADDITIONAL INPUTS ------------------
st.subheader("Step 2: Select City Tier")
city_tier = st.selectbox(
    "City Tier",
    ["Tier 1 (Metros)", "Tier 2 (Emerging Cities)", "Tier 3 (Small Cities)", "Tier 4 (Rural/Remote)"]
)

st.subheader("Step 3: Optional Role/Department Tag")
role_tag = st.text_input(
    "Role / Department (optional):",
    placeholder="e.g., Project Lead / Manager / Analyst"
).strip()

st.subheader("Step 4: Enter Comparator Organizations")
user_input = st.text_input(
    "Enter comparator organizations (comma-separated):",
    placeholder="e.g., Teach For India, Pratham, Dalberg, Samagra"
)
comparators = [x.strip() for x in user_input.split(",") if x.strip()]

st.subheader("Step 5: Playbook Depth")
depth_option = st.radio(
    "Select depth of playbook:",
    ["High-Level Summary", "Detailed Deep-Dive"],
    index=1
)

# ------------------ MASTER PROMPT ------------------
MASTER_PROMPT = f"""
You are an expert Talent Acquisition Strategist, Labor Market Economist, 
and Organizational Psychologist. Create a **consulting-style implementation playbook**
for talent acquisition based on the following inputs.

### Inputs:
- Job Description: {jd_text if jd_text else 'Not provided'}
- City Tier: {city_tier}
- Role/Department: {role_tag if role_tag else 'Not specified'}
- Comparator Organizations: {', '.join(comparators) if comparators else 'None provided'}
- Playbook Depth: {depth_option}

### Deliverables:
1. Market scan & sourcing opportunities for this role and city tier.
2. Comparator benchmarking scorecard:
   - Sourcing channels
   - Compensation & benefits
   - Retention & turnover patterns
   - Employer branding & positioning
   - Learning & skill development programs
3. Candidate personas (motivators, barriers, skill gaps).
4. Sourcing channel analysis (traditional vs non-traditional), ROI, funnel effectiveness, costs.
5. Implementation roadmap with milestones and sequencing.
6. Risk assessment & mitigation strategies (e.g., relocation challenges).
7. KPIs: qualified leads, funnel conversion, time-to-hire, cost-per-hire, % hires via attitude/agility track.
8. Recommendations: balancing experience vs attitude/agility.
9. Candidate interview guidance: 5–7 structured interviews per city tier.
10. Decision gates: scale/don’t scale based on blended quantitative and qualitative metrics.
11. JD Flexibility: playbook should be reusable for different roles.
"""

# ------------------ GENERATE PLAYBOOK ------------------
if st.button("🚀 Generate Playbook"):
    if not api_key:
        st.warning("Please enter your OpenAI API Key first.")
    elif not jd_text or len(jd_text) < 20:
        st.warning("Please upload a valid JD (at least 20 characters).")
    else:
        with st.spinner("Generating playbook... this may take a moment ⏳"):
            try:
                response = client.chat.completions.create(
                    model="gpt-5",
                    messages=[
                        {"role": "system", "content": "You are a management consultant and TA strategist."},
                        {"role": "user", "content": MASTER_PROMPT},
                    ]
                )
                output = response.choices[0].message.content
                st.success("✅ Playbook generated!")

                # ------------------ DISPLAY FULL PLAYBOOK ------------------
                st.markdown("### 📄 Talent Acquisition Playbook")
                st.text_area("Full Playbook", output, height=600)

                # ------------------ COPY BUTTON ------------------
                st.button("📋 Copy Full Playbook to Clipboard", on_click=lambda: st.session_state.update({"clipboard": output}))

                # ------------------ DOWNLOAD OPTIONS ------------------
                # 1. Download as TXT
                st.download_button(
                    label="💾 Download Playbook as TXT",
                    data=output,
                    file_name="TA_Playbook.txt",
                    mime="text/plain"
                )

                # 2. Download as Word (.docx) with formatting
                doc = Document()
                doc.add_heading("Talent Acquisition Playbook", 0)
                for line in output.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    # Bold section headers (lines starting with number + dot)
                    if line[0].isdigit() and "." in line[:4]:
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                        run.font.size = Pt(12)
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.font.size = Pt(11)
                        # Italicize bullets starting with '-'
                        if line.startswith("-"):
                            run.italic = True
                doc_io = BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                st.download_button(
                    label="💾 Download Playbook as Word (.docx)",
                    data=doc_io,
                    file_name="TA_Playbook.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                # 3. Download as PDF with formatting
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=12)

                # Replace unsupported chars for latin-1
                safe_output = output.encode('latin-1', errors='replace').decode('latin-1')
                for line in safe_output.split("\n"):
                    line = line.strip()
                    if not line:
                        pdf.ln(4)
                        continue
                    # Bold section headers
                    if line[0].isdigit() and "." in line[:4]:
                        pdf.set_font("Arial", style="B", size=12)
                    # Italics for bullets
                    elif line.startswith("-"):
                        pdf.set_font("Arial", style="I", size=12)
                    else:
                        pdf.set_font("Arial", style="", size=12)
                    pdf.multi_cell(0, 8, line)

                pdf_io = BytesIO()
                pdf_bytes = pdf.output(dest='S').encode('latin-1')
                pdf_io.write(pdf_bytes)
                pdf_io.seek(0)
                st.download_button(
                    label="💾 Download Playbook as PDF",
                    data=pdf_io,
                    file_name="TA_Playbook.pdf",
                    mime="application/pdf"
                )

            except Exception as e:
                st.error(f"Error generating playbook: {e}")
